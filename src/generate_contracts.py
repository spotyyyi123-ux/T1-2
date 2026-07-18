"""
Генератор синтетических договоров для датасета Sprint 1.
Создаёт DOCX-файлы и эталонную разметку labels.json.

Запуск:
    python src/generate_contracts.py --count 30 --out ./dataset
"""

import argparse
import json
import random
from datetime import date, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --------------------------------------------------------------------------
# СПРАВОЧНИКИ ДАННЫХ
# --------------------------------------------------------------------------

CITIES = [
    "Москва", "Санкт-Петербург", "Новосибирск", "Екатеринбург",
    "Казань", "Нижний Новгород", "Челябинск", "Самара", "Омск",
    "Ростов-на-Дону", "Уфа", "Красноярск", "Воронеж", "Пермь", "Волгоград"
]

COMPANY_FORMS = ["ООО", "АО", "ПАО", "ИП"]

COMPANY_NAMES = [
    "Ромашка", "Вертикаль", "СтройГрад", "ТехноЛайн", "СевернаяЗвезда",
    "БелаяБаза", "ГлобалТрейд", "ИнтерСервис", "МаксимаГрупп", "ЮгСтрой",
    "АльфаТех", "СибирьПром", "ГазИнвест", "ЭнергоПроект", "ИТ-Решения",
    "АгроТрейд", "МеталлИндустрия", "ХимКомплект", "СвязьМонтаж", "АвтоЛига"
]

STREETS = [
    "Ленина", "Мира", "Пушкина", "Гагарина", "Советская",
    "Октябрьская", "Молодёжная", "Садовая", "Набережная", "Лесная"
]

SURNAMES = ["Иванов", "Петров", "Сидоров", "Кузнецов", "Смирнов",
            "Попов", "Васильев", "Михайлов", "Новиков", "Фёдоров"]

NAMES = ["Иван", "Пётр", "Александр", "Дмитрий", "Сергей",
         "Алексей", "Андрей", "Максим", "Михаил", "Николай"]

CONTRACT_KINDS = [
    "поставки товара",
    "оказания услуг",
    "аренды нежилого помещения",
    "подряда на выполнение работ",
    "возмездного оказания консультационных услуг",
    "выполнения проектных работ",
]

MONTHS = ["января", "февраля", "марта", "апреля", "мая", "июня",
          "июля", "августа", "сентября", "октября", "ноября", "декабря"]

# --------------------------------------------------------------------------
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# --------------------------------------------------------------------------

def number_to_words_ru(n: int) -> str:
    """Конвертирует число до 999 999 в пропись."""
    if n == 0:
        return "ноль"
    
    units = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = ["десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
             "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать"]
    tens = ["", "", "двадцать", "тридцать", "сорок", "пятьдесят",
            "шестьдесят", "семьдесят", "восемьдесят", "девяносто"]
    hundreds = ["", "сто", "двести", "триста", "четыреста",
                "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот"]
    
    def convert_hundreds(num):
        if num == 0:
            return ""
        result = []
        h = num // 100
        if h:
            result.append(hundreds[h])
        remainder = num % 100
        if remainder:
            if 10 <= remainder <= 19:
                result.append(teens[remainder - 10])
            else:
                t = remainder // 10
                u = remainder % 10
                if t:
                    result.append(tens[t])
                if u:
                    result.append(units[u])
        return " ".join(result)
    
    thousands = n // 1000
    rest = n % 1000
    
    parts = []
    if thousands:
        th_text = convert_hundreds(thousands)
        if thousands % 10 == 1 and thousands % 100 != 11:
            parts.append(f"{th_text} тысяча")
        elif 2 <= thousands % 10 <= 4 and not (12 <= thousands % 100 <= 14):
            parts.append(f"{th_text} тысячи")
        else:
            parts.append(f"{th_text} тысяч")
    
    if rest:
        parts.append(convert_hundreds(rest))
    
    return " ".join(parts)


def generate_contract_number() -> str:
    """Генерирует номер договора в разных форматах."""
    formats = [
        f"№ {random.randint(1, 999)}-{random.choice(['Д', 'У', 'П', 'С'])}/{random.randint(2022, 2024)}",
        f"№ {random.randint(1, 999)}/{random.randint(2022, 2024)}",
        f"{random.randint(1, 999)}-{random.choice(['Д', 'У', 'П'])}",
    ]
    return random.choice(formats)


def generate_date() -> date:
    """Генерирует случайную дату в диапазоне 2022-2024."""
    start = date(2022, 1, 1)
    days_range = (date(2024, 12, 31) - start).days
    return start + timedelta(days=random.randint(0, days_range))


def format_date(d: date, style: int = None) -> str:
    """Форматирует дату в разных стилях."""
    styles = [
        f"{d.day:02d}.{d.month:02d}.{d.year}",
        f"{d.day:02d} {MONTHS[d.month-1]} {d.year} г.",
        f"«{d.day:02d}» {MONTHS[d.month-1]} {d.year} г.",
        f"{d.year}-{d.month:02d}-{d.day:02d}",
    ]
    if style is not None:
        return styles[style]
    return random.choice(styles)


def generate_company() -> dict:
    """Генерирует компанию со всеми реквизитами."""
    form = random.choice(COMPANY_FORMS)
    name = random.choice(COMPANY_NAMES)
    full_name = f'{form} «{name}»'
    
    if form == "ИП":
        inn = ''.join(random.choices('0123456789', k=12))
    else:
        inn = ''.join(random.choices('0123456789', k=10))
    
    ogrn = ''.join(random.choices('0123456789', k=13))
    
    return {
        "full_name": full_name,
        "inn": inn,
        "ogrn": ogrn,
        "address": f"{random.randint(100000, 999999)}, г. {random.choice(CITIES)}, ул. {random.choice(STREETS)}, д. {random.randint(1, 150)}",
        "signer_surname": random.choice(SURNAMES),
        "signer_name": random.choice(NAMES),
    }


def generate_amount() -> dict:
    """Генерирует случайную сумму."""
    amount = random.randint(50_000, 950_000)
    if random.random() > 0.3:
        amount = (amount // 1000) * 1000
    
    return {
        "numeric": amount,
        "formatted": f"{amount:,}".replace(",", " "),
        "words": number_to_words_ru(amount) + " рублей 00 копеек",
    }


# --------------------------------------------------------------------------
# ШАБЛОНЫ ДОГОВОРОВ
# --------------------------------------------------------------------------

TEMPLATES = [
    {
        "name": "classic",
        "text": """ДОГОВОР {{KIND}}
{{NUMBER}}

{{CITY}}                                                       {{DATE}}

{{CUSTOMER_FULL}}, именуемое в дальнейшем «Заказчик», в лице
генерального директора {{CUSTOMER_SIGNER_SURNAME}} {{CUSTOMER_SIGNER_NAME}},
действующего на основании Устава, с одной стороны, и
{{EXECUTOR_FULL}}, именуемое в дальнейшем «Исполнитель», в лице
директора {{EXECUTOR_SIGNER_SURNAME}} {{EXECUTOR_SIGNER_NAME}},
действующего на основании Устава, с другой стороны, совместно
именуемые «Стороны», заключили настоящий Договор о нижеследующем.

1. ПРЕДМЕТ ДОГОВОРА
1.1. Исполнитель обязуется оказать услуги, а Заказчик обязуется принять
и оплатить их в порядке и на условиях, предусмотренных настоящим Договором.

2. ЦЕНА ДОГОВОРА
2.1. Общая стоимость услуг составляет {{AMOUNT_FORMATTED}} ({{AMOUNT_WORDS}}).
2.2. Оплата производится в течение 10 рабочих дней после подписания акта.

3. РЕКВИЗИТЫ СТОРОН

ЗАКАЗЧИК:
{{CUSTOMER_FULL}}
ИНН {{CUSTOMER_INN}}, ОГРН {{CUSTOMER_OGRN}}
Адрес: {{CUSTOMER_ADDRESS}}

ИСПОЛНИТЕЛЬ:
{{EXECUTOR_FULL}}
ИНН {{EXECUTOR_INN}}, ОГРН {{EXECUTOR_OGRN}}
Адрес: {{EXECUTOR_ADDRESS}}

_____________________ /{{CUSTOMER_SIGNER_SURNAME}}/
_____________________ /{{EXECUTOR_SIGNER_SURNAME}}/
М.П.                                   М.П.
"""
    },
    
    {
        "name": "compact",
        "text": """ДОГОВОР {{KIND}} {{NUMBER}}

{{CITY}}                              от {{DATE}}

{{CUSTOMER_FULL}} (далее — «Сторона 1») и {{EXECUTOR_FULL}}
(далее — «Сторона 2») заключили настоящий договор:

1. Сторона 2 обязуется выполнить работы по заданию Стороны 1.
2. Цена договора: {{AMOUNT_FORMATTED}} руб. ({{AMOUNT_WORDS}}).
3. Настоящий договор {{NUMBER}} составлен в двух экземплярах,
имеющих равную юридическую силу.

СТОРОНЫ ДОГОВОРА:

Сторона 1: {{CUSTOMER_FULL}}, ИНН {{CUSTOMER_INN}}
Сторона 2: {{EXECUTOR_FULL}}, ИНН {{EXECUTOR_INN}}

______________________                ______________________
{{CUSTOMER_SIGNER_SURNAME}}           {{EXECUTOR_SIGNER_SURNAME}}
"""
    },
    
    {
        "name": "split_amount",
        "text": """ДОГОВОР {{KIND}}
{{NUMBER}} от {{DATE}}

Мы, нижеподписавшиеся:

{{CUSTOMER_FULL}}, в лице {{CUSTOMER_SIGNER_SURNAME}} {{CUSTOMER_SIGNER_NAME}},
именуемый далее «Заказчик», и

{{EXECUTOR_FULL}}, в лице {{EXECUTOR_SIGNER_SURNAME}} {{EXECUTOR_SIGNER_NAME}},
именуемый далее «Подрядчик»,

договорились о следующем:

1. Подрядчик выполняет работы согласно техническому заданию.
2. Стоимость работ составляет {{AMOUNT_FORMATTED}} рублей 00 копеек.
3. Указанная сумма в размере {{AMOUNT_FORMATTED}} руб. уплачивается
Заказчиком в следующем порядке: 50% предоплата, 50% по факту.
4. Прописью сумма договора: {{AMOUNT_WORDS}}.

Реквизиты:

Заказчик: {{CUSTOMER_FULL}}, {{CUSTOMER_ADDRESS}}, ИНН {{CUSTOMER_INN}}
Подрядчик: {{EXECUTOR_FULL}}, {{EXECUTOR_ADDRESS}}, ИНН {{EXECUTOR_INN}}
"""
    },
    
    {
        "name": "framework",
        "text": """РАМОЧНЫЙ ДОГОВОР {{KIND}} {{NUMBER}}

г. {{CITY}}                                     {{DATE}}

{{CUSTOMER_FULL}} (Заказчик) и {{EXECUTOR_FULL}} (Исполнитель)
заключили настоящий рамочный договор о нижеследующем:

Статья 1. Предмет
Исполнитель обязуется по заявкам Заказчика оказывать услуги,
перечень и стоимость которых согласовываются в Приложениях.

Статья 2. Цена
Максимальная сумма договора: {{AMOUNT_FORMATTED}} руб.
({{AMOUNT_WORDS}}).

Статья 3. Заявки
Каждая заявка оформляется отдельным Приложением к настоящему
Договору {{NUMBER}}.

Статья 4. Реквизиты
Заказчик: {{CUSTOMER_FULL}}, ИНН {{CUSTOMER_INN}}, ОГРН {{CUSTOMER_OGRN}}
Исполнитель: {{EXECUTOR_FULL}}, ИНН {{EXECUTOR_INN}}, ОГРН {{EXECUTOR_OGRN}}

От Заказчика:                От Исполнителя:
________/{{CUSTOMER_SIGNER_SURNAME}}/     ________/{{EXECUTOR_SIGNER_SURNAME}}/
"""
    },
    
    {
        "name": "modern",
        "text": """ДОГОВОР {{KIND}} № {{NUMBER_SHORT}}

от {{DATE_DIGITAL}}

{{CUSTOMER_FULL}} и {{EXECUTOR_FULL}}, вместе именуемые «Стороны»,
заключили настоящий договор на следующих условиях:

1. Услуги оказываются в соответствии с Договором № {{NUMBER_SHORT}}.
2. Общая цена: {{AMOUNT_FORMATTED}} руб. ({{AMOUNT_WORDS}}).
3. Договор № {{NUMBER_SHORT}} вступает в силу с {{DATE_DIGITAL}}.

Подписи Сторон:

{{CUSTOMER_FULL}}            {{EXECUTOR_FULL}}
________________             ________________
({{CUSTOMER_SIGNER_SURNAME}})           ({{EXECUTOR_SIGNER_SURNAME}})
"""
    },
]


# --------------------------------------------------------------------------
# ГЕНЕРАЦИЯ ОДНОГО ДОГОВОРА
# --------------------------------------------------------------------------

def generate_one_contract(index: int) -> dict:
    """Генерирует один договор и возвращает текст + эталонную разметку."""
    template = random.choice(TEMPLATES)
    
    customer = generate_company()
    executor = generate_company()
    while executor["full_name"] == customer["full_name"]:
        executor = generate_company()
    
    contract_number = generate_contract_number()
    number_short = f"{random.randint(1, 999)}"
    contract_date = generate_date()
    amount = generate_amount()
    contract_kind = random.choice(CONTRACT_KINDS)
    city = random.choice(CITIES)
    
    text = template["text"]
    text = text.replace("{{KIND}}", contract_kind)
    text = text.replace("{{NUMBER}}", contract_number)
    text = text.replace("{{NUMBER_SHORT}}", number_short)
    text = text.replace("{{CITY}}", city)
    text = text.replace("{{DATE}}", format_date(contract_date))
    text = text.replace("{{DATE_DIGITAL}}", format_date(contract_date, style=0))
    
    text = text.replace("{{CUSTOMER_FULL}}", customer["full_name"])
    text = text.replace("{{CUSTOMER_INN}}", customer["inn"])
    text = text.replace("{{CUSTOMER_OGRN}}", customer["ogrn"])
    text = text.replace("{{CUSTOMER_ADDRESS}}", customer["address"])
    text = text.replace("{{CUSTOMER_SIGNER_SURNAME}}", customer["signer_surname"])
    text = text.replace("{{CUSTOMER_SIGNER_NAME}}", customer["signer_name"])
    
    text = text.replace("{{EXECUTOR_FULL}}", executor["full_name"])
    text = text.replace("{{EXECUTOR_INN}}", executor["inn"])
    text = text.replace("{{EXECUTOR_OGRN}}", executor["ogrn"])
    text = text.replace("{{EXECUTOR_ADDRESS}}", executor["address"])
    text = text.replace("{{EXECUTOR_SIGNER_SURNAME}}", executor["signer_surname"])
    text = text.replace("{{EXECUTOR_SIGNER_NAME}}", executor["signer_name"])
    
    text = text.replace("{{AMOUNT_FORMATTED}}", amount["formatted"])
    text = text.replace("{{AMOUNT_WORDS}}", amount["words"])
    
    label = {
        "file": f"contract_{index:04d}.docx",
        "template": template["name"],
        "number": contract_number,
        "date": contract_date.isoformat(),
        "date_text": format_date(contract_date),
        "amount": amount["numeric"],
        "amount_text": amount["words"],
        "kind": contract_kind,
        "parties": [
            {
                "name": customer["full_name"],
                "role": "Заказчик",
                "inn": customer["inn"],
                "ogrn": customer["ogrn"],
            },
            {
                "name": executor["full_name"],
                "role": "Исполнитель",
                "inn": executor["inn"],
                "ogrn": executor["ogrn"],
            }
        ]
    }
    
    return {"text": text, "label": label}


def save_docx(text: str, filepath: Path):
    """Сохраняет текст в форматированный DOCX."""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for para_text in text.strip().split('\n'):
        if para_text.startswith('ДОГОВОР') or para_text.startswith('РАМОЧНЫЙ'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(para_text)
            run.bold = True
            run.font.size = Pt(14)
        else:
            p = doc.add_paragraph(para_text)
    
    doc.save(str(filepath))


def main():
    parser = argparse.ArgumentParser(description="Генератор синтетических договоров")
    parser.add_argument("--count", type=int, default=30, help="Количество договоров")
    parser.add_argument("--out", type=str, default="./dataset", help="Выходная папка")
    parser.add_argument("--seed", type=int, default=42, help="Seed для воспроизводимости")
    args = parser.parse_args()
    
    random.seed(args.seed)
    
    out_dir = Path(args.out)
    docx_dir = out_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    
    print(f"Генерация {args.count} договоров...")
    print(f"Папка: {docx_dir.resolve()}")
    print()
    
    all_labels = []
    
    for i in range(1, args.count + 1):
        contract = generate_one_contract(i)
        filename = f"contract_{i:04d}.docx"
        save_docx(contract["text"], docx_dir / filename)
        all_labels.append(contract["label"])
        
        if i % 5 == 0:
            print(f"  Готово {i}/{args.count}...")
    
    labels_dir = out_dir / "annotations"
    labels_dir.mkdir(parents=True, exist_ok=True)
    labels_path = labels_dir / "labels.json"
    
    with open(labels_path, "w", encoding="utf-8") as f:
        json.dump(all_labels, f, ensure_ascii=False, indent=2)
    
    print(f"\nГотово! Создано {args.count} договоров")
    print(f"DOCX: {docx_dir.resolve()}")
    print(f"Разметка: {labels_path.resolve()}")
    
    print("\nСтатистика по шаблонам:")
    templates_used = {}
    for label in all_labels:
        t = label["template"]
        templates_used[t] = templates_used.get(t, 0) + 1
    for t, c in templates_used.items():
        print(f"  {t}: {c} шт.")


if __name__ == "__main__":
    main()
